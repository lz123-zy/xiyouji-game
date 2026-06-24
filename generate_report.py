# -*- coding: utf-8 -*-
"""生成西游记小游戏项目需求分析报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '黑体'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.size = Pt(18 - level * 2)

# ── 辅助函数 ──
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_info(label, value):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(value)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    return table

def add_flow(lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.2
        for r in p.runs:
            r.font.size = Pt(10.5)
            r.font.name = 'Consolas'

# ═══════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title('课程学习小组讨论记录')
doc.add_paragraph()
add_info('【讨论主题】：', '项目需求分析及设计方案——讨论记录')
add_info('【课程名称】：', '软件项目开发与实践')
add_info('【项目名称】：', '西游记·祸起观音院（2D RPG小游戏）')
add_info('【讨论日期】：', '2026年6月24日')
doc.add_page_break()

# ═══════════════════════════════════════════════════
# 一、项目背景与目标
# ═══════════════════════════════════════════════════
doc.add_heading('一、项目背景与目标', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '本项目以中国古典名著《西游记》为题材，开发一款基于 Python + Pygame 的 2D 俯视角 RPG 小游戏。'
    '玩家操控孙悟空（齐天大圣），在村庄、郊外、观音院三个场景中探索，与 NPC 对话获取任务，'
    '击败妖怪，最终击败 Boss（牛魔王）完成冒险。'
)
doc.add_paragraph(
    '当前市面上以《西游记》为主题的独立小游戏较少，且多为卡牌或回合制，'
    '缺乏沉浸式探索+即时战斗的体验。本项目旨在用轻量级技术栈实现一个完整的 RPG 小游戏原型，'
    '可作为 Pygame 2D RPG 开发的学习参考案例。'
)

doc.add_heading('1.2 项目目标', level=2)

doc.add_heading('功能性目标', level=3)
func_goals = [
    '实现三场景（村庄 → 郊外 → 观音院）的无缝切换与探索系统',
    '实现 NPC 对话系统，支持按角色、按任务阶段显示不同对话',
    '实现完整的主线任务流程（接取 → 战斗 → 击败Boss → 复命）',
    '实现怪物 AI 系统（巡逻 → 追击 → 返回）与回合制战斗系统',
    '实现 Boss 战（牛魔王），含狂暴阶段（50% 血量触发）',
    '实现角色四向移动动画、碰撞检测、镜头跟随',
]
for g in func_goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('非功能性目标', level=3)
nonfunc_goals = [
    '性能：60 FPS 稳定运行，800×600 窗口分辨率',
    '可用性：WASD/方向键移动，E/空格交互，J 攻击，操作简洁直观',
    '容错性：所有资源加载均有 try/except 兜底，缺失素材自动降级为占位图',
    '可扩展性：地图使用 Tiled（TMX）格式，NPC/怪物均通过地图对象层配置，便于内容扩展',
]
for g in nonfunc_goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 二、主要讨论内容
# ═══════════════════════════════════════════════════
doc.add_heading('二、主要讨论内容', level=1)

# ── 2.1 功能需求分析 ──
doc.add_heading('2.1 功能需求分析', level=2)

add_table(
    ['编号', '功能点', '说明', '优先级'],
    [
        ['F01', '角色移动与动画',
         '孙悟空支持 WASD/方向键四向移动，每方向4帧动画，支持斜向移动（速度归一化），碰撞后支持沿墙滑动', '高'],
        ['F02', '镜头跟随系统',
         '摄像机以玩家为中心跟随，自动限制在地图边界内', '高'],
        ['F03', 'NPC对话系统',
         '按E/空格与NPC交互，底部弹出对话框显示角色名+台词；支持按对象名（个体专属）和图层名（类型兜底）两级解析', '高'],
        ['F04', '任务状态机',
         '五阶段状态机：NOT_ACCEPTED→ACCEPTED→CLEARED→RETURNED→COMPLETE，NPC对话随阶段变化', '高'],
        ['F05', '三场景切换',
         '村庄（探索+对话）→ 郊外（过渡+NPC）→ 观音院（战斗），返回流程反向；每个场景独立加载、独立碰撞', '高'],
        ['F06', '怪物AI系统',
         '巡逻（随机走动）→ 追击（玩家进入感知范围）→ 返回（玩家离开后回家），含速度/距离/状态机管理', '高'],
        ['F07', '回合制战斗',
         '暗色遮罩+战斗面板，玩家按J攻击（20伤害），怪物自动攻击（12伤害），双方独立冷却；含攻击动画、受击闪红、魔法特效', '高'],
        ['F08', 'Boss战',
         '牛魔王（160HP），50%血量进入狂暴（攻击24，冷却1.0s），Boss出场横幅公告', '高'],
        ['F09', '胜利/失败处理',
         '胜利后显示胜利画面，失败后可按R重试当前怪物，按Esc退出', '高'],
        ['F10', '开始界面',
         '程序化绘制夜空背景（星星、月亮、山脉），显示标题、副标题、剧情简介、操作提示', '中'],
        ['F11', '帮助/暂停界面',
         '按H查看操作说明，按P暂停（保留状态可恢复）', '中'],
        ['F12', '音频管理',
         '背景音乐循环播放，攻击音效，胜利音效，按M全局静音', '中'],
        ['F13', '全屏切换',
         'F11或Alt+Enter切换全屏/窗口模式', '低'],
        ['F14', '控制提示HUD',
         '左上角显示当前可用操作键位和音频/全屏状态', '低'],
        ['F15', '怪物消失特效',
         '怪物死亡后播放魔法消失动画，再从场景中移除', '中'],
    ]
)

doc.add_paragraph()

# ── 2.2 技术方案设计 ──
doc.add_heading('2.2 技术方案设计', level=2)

add_table(
    ['层级', '技术选型', '说明'],
    [
        ['编程语言', 'Python 3.13', '开发效率高，生态丰富，适合原型开发'],
        ['游戏引擎', 'Pygame 2.6.1', '轻量级2D游戏库，适合本项目规模；无额外引擎依赖'],
        ['地图编辑', 'Tiled Editor + pytmx 3.32', '业界标准TMX地图格式，支持对象层/碰撞层/图层；pytmx负责运行时加载'],
        ['精灵素材', 'TGA/PNG序列帧', '角色/怪物/NPC使用TGA格式序列帧动画，UI使用PNG，场景使用JPG大图'],
        ['精灵处理', '自研process_npc.py', '基于直方图的色度键抠图，支持逐角色调参的亮度/对比度/饱和度/RGB缩放'],
        ['音频', 'Pygame mixer', '内建音频模块，支持BGM循环+音效播放，静音开关'],
        ['地图方案', '照片级背景+TMX碰撞层', '村庄/观音院使用实景照片做背景，碰撞数据通过TMX对象层矩形障碍物定义'],
    ]
)

doc.add_paragraph()

# ── 2.3 核心业务流程 ──
doc.add_heading('2.3 核心业务流程', level=2)

doc.add_paragraph('主流程：')
add_flow([
    '[开始画面] ──按Enter──→ [村庄探索]',
    '    │',
    '    ├─ 与土地公对话 → 接受任务（NOT_ACCEPTED → ACCEPTED）',
    '    │',
    '    ├──按E/空格──→ [郊外探索]',
    '    │                   │',
    '    │                   ├──走到右侧出口──→ [观音院]',
    '    │                   │                     │',
    '    │                   │                     ├─ 击败小怪 ×4',
    '    │                   │                     ├─ Boss牛魔王出现',
    '    │                   │                     ├─ 击败Boss → CLEARED',
    '    │                   │                     │',
    '    │                   │←──自动传送──← [返回郊外]',
    '    │                   │',
    '    │←──走到左侧出口──← [返回村庄]',
    '    │',
    '    ├─ 与土地公复命 → COMPLETE',
    '    └──→ [胜利画面] ──按Esc──→ [退出]',
])

doc.add_paragraph()
doc.add_paragraph('战斗子流程：')
add_flow([
    '玩家触发怪物 trigger_rect',
    '    │',
    '    ├──→ [战斗面板]',
    '    │       ├─ 玩家按J → 播放攻击动画 → 扣除怪物20 HP',
    '    │       ├─ 怪物自动攻击 → 扣除玩家12~24 HP',
    '    │       ├─ 怪物HP ≤ 0 → 胜利 → 返回探索',
    '    │       └─ 玩家HP ≤ 0 → 失败 → 可按R重试',
    '    │',
    '    └──→ [所有小怪击败] → Boss出场横幅 → Boss战斗',
])

doc.add_page_break()

# ── 2.4 地图与场景设计 ──
doc.add_heading('2.4 地图与场景设计', level=2)

doc.add_heading('场景概览', level=3)
add_table(
    ['场景', '地图文件', '尺寸（像素）', '实体', '特点'],
    [
        ['村庄', 'village1.tmx', '3780×2395', '土地公×1、老人×3',
         '照片背景+148个碰撞矩形；NPC对话推动任务'],
        ['郊外', 'scene.tmx / 郊外.jpg', '1376×768', '守林人×1、采药人×1、草药娘×1、商人×1',
         '传统瓦片地图或照片背景；过渡区域，左右出口'],
        ['观音院', 'temple1.tmx', '1999×1495', '小怪×4 + Boss×1',
         '照片背景+139个碰撞矩形；怪物AI巡逻区'],
    ]
)

doc.add_paragraph()
doc.add_heading('NPC角色列表（当前版本）', level=3)
add_table(
    ['NPC ID', '角色名', '所在场景', '对话示例'],
    [
        ['god', '土地公', '村庄', '"大圣，前方就是观音院，小心妖怪。"'],
        ['elder1', '张老汉', '村庄', '"听闻观音院的禅师都被妖怪掳走了，大圣可要救救他们啊。"'],
        ['elder2', '王婆婆', '村庄', '"老身年轻时也见过妖怪作乱，那回还是一位高僧把妖降住的。"'],
        ['elder4', '赵老伯', '村庄', '"大圣若要进观音院，记得先找土地公问个明白。"'],
        ['guard', '老林', '郊外', '"我是这林子的守林人。前方妖怪横行，大圣务必多加小心。"'],
        ['apprentice', '小陆', '郊外', '"我跟着师父学采药，听说观音院的药园里有一味灵芝..."'],
        ['herbalist', '芸娘', '郊外', '"大圣若在林间受伤，可到我这儿来讨些草药。"'],
        ['merchant', '沈老板', '郊外', '"这妖怪一闹，我这生意也做不成了。大圣若能除掉那黑熊精，我请全村人喝酒！"'],
    ]
)

doc.add_page_break()

# ── 2.5 主要难点 ──
doc.add_heading('2.5 主要难点与解决思路', level=2)

add_table(
    ['难点', '说明', '解决思路'],
    [
        ['照片级地图的碰撞设计',
         '村庄/观音院使用实景照片做背景，无法像瓦片地图那样自动生成碰撞；需手动在Tiled中放置100+个矩形碰撞框',
         '使用Tiled编辑器的对象层手动绘制碰撞矩形；已在村庄放置148个、观音院放置139个障碍物'],
        ['NPC素材来源不统一',
         '土地公/老人有现成TGA序列帧；郊外NPC需从游戏截图中抠图生成',
         '自研process_npc.py脚本：基于直方图白色背景去除+亮度/对比度调色+边缘裁切，批量生成透明背景PNG'],
        ['怪物AI状态切换的平滑性',
         '巡逻→追击→返回三状态切换时，需避免怪物"卡住"或"瞬移"',
         '采用状态机+每状态独立的速度/计时器；追击有距离上限（leash=260px）防止跑出地图；返回时缓慢走回'],
        ['Boss狂暴阶段的平衡性',
         '需要在不增加代码复杂度的前提下让Boss有明显难度提升',
         '在Boss类中通过属性覆盖实现：HP<50%时切换攻击力(16→24)和冷却(1.5→1.0s)，视觉上有状态文字提示'],
        ['场景切换的连续性',
         '三个场景的玩家出生点、任务状态、NPC列表需在切换时正确保持',
         '每次切换重建Scene和Player对象，通过QuestManager保持全局任务状态；出生点按方向设定'],
    ]
)

doc.add_page_break()

# ── 2.6 项目文件结构 ──
doc.add_heading('2.6 项目文件结构', level=2)

add_flow([
    '西游记小游戏/',
    '├── main.py                      # 入口文件',
    '├── requirements.txt             # 依赖：pygame, pytmx',
    '├── generate_art.py              # 程序化场景素材生成器',
    '├── process_npc.py               # NPC截图抠图处理工具',
    '├── 项目需求分析报告.docx          # 本文档',
    '├── src/                         # 源代码目录（18个Python文件）',
    '│   ├── game.py                  # 主游戏控制器（634行）',
    '│   ├── game_state.py            # 游戏状态枚举（10种状态）',
    '│   ├── player.py                # 玩家角色（四向动画+碰撞）',
    '│   ├── npc.py                   # NPC实体（对话+动画）',
    '│   ├── monster.py               # 怪物AI（巡逻/追击/返回）',
    '│   ├── boss.py                  # Boss子类（牛魔王+狂暴）',
    '│   ├── battle.py                # 回合制战斗系统',
    '│   ├── quest.py                 # 任务状态机（5阶段）',
    '│   ├── scene.py                 # 场景管理器',
    '│   ├── tmx_map.py               # TMX地图加载器',
    '│   ├── camera.py                # 摄像机跟随系统',
    '│   ├── dialog.py                # 对话框渲染器',
    '│   ├── ui.py                    # UI界面系统（310行）',
    '│   ├── audio.py                 # 音频管理器',
    '│   ├── animation.py             # 帧动画控制器',
    '│   └── settings.py              # 全局配置常量',
    '├── data/                        # 数据配置',
    '│   ├── dialogs.json             # 对话树定义（18条）',
    '│   ├── monsters.json            # 怪物属性定义',
    '│   └── scenes.json              # 场景元数据',
    '├── resource/                    # 游戏资源（621个文件）',
    '│   ├── tmx/                     # Tiled地图文件（6个TMX）',
    '│   ├── img/                     # 图片素材（角色/怪物/UI/场景）',
    '│   ├── font/                    # 中文字体',
    '│   └── sound/                   # 音频文件（BGM+音效）',
    '└── NPC形象/                     # NPC原始截图素材',
])

doc.add_page_break()

# ── 2.7 模块依赖关系 ──
doc.add_heading('2.7 模块依赖关系', level=2)

add_flow([
    'main.py → Game',
    '          ├── settings（全局常量）',
    '          ├── game_state（状态枚举）',
    '          ├── quest（任务管理）',
    '          ├── player（玩家）',
    '          │     └── settings',
    '          ├── scene（场景管理）',
    '          │     ├── tmx_map（地图加载）',
    '          │     ├── npc（NPC实体）',
    '          │     │     ├── animation',
    '          │     │     └── settings',
    '          │     └── monster（怪物AI）',
    '          │           ├── animation',
    '          │           └── settings',
    '          ├── boss（Boss，继承Monster）',
    '          ├── battle（战斗系统）',
    '          │     ├── animation',
    '          │     └── settings',
    '          ├── camera（摄像机）',
    '          ├── dialog（对话框）',
    '          │     └── settings',
    '          ├── ui（界面渲染）',
    '          └── audio（音频管理）',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 三、总结
# ═══════════════════════════════════════════════════
doc.add_heading('三、总结', level=1)
doc.add_paragraph(
    '本项目采用 Python + Pygame 技术栈，以《西游记》观音院故事为背景，'
    '实现了一个包含探索、对话、任务、战斗的完整 2D RPG 小游戏原型。'
    '核心架构为模块化的 Game 控制器 + Scene 场景管理 + QuestManager 任务状态机，'
    '通过 Tiled 编辑器的地图对象层驱动 NPC 和怪物的配置，实现了数据与逻辑的分离。'
)
doc.add_paragraph(
    '项目共包含 18 个 Python 源文件（约 2800 行代码）、3 个 JSON 数据文件、'
    '6 个 TMX 地图文件和 621 个游戏资源文件。'
    '通过本项目的开发实践，完整体验了从需求分析、技术选型、架构设计到编码实现的软件开发全流程。'
)

# ── 保存 ──
output = r'C:\Users\zhx\Desktop\西游记小游戏\项目需求分析报告.docx'
doc.save(output)
print(f'报告已生成：{output}')
